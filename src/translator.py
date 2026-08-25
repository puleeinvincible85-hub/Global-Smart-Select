from io import BytesIO

from docx import Document

from src.ai_extractor import get_openai_client


def translate_document(
    document_text,
    target_language,
    standard_mappings=None,
):
    """
    Translate document text into the selected language.

    If approved standard mappings are supplied,
    the translation will use those enterprise-standard
    terms where appropriate.
    """

    client = get_openai_client()

    mappings_text = ""

    if standard_mappings:
        mapping_lines = []

        for mapping in standard_mappings:
            source_value = mapping.get("source_value")
            canonical_value = mapping.get("canonical_value")

            if source_value and canonical_value:
                mapping_lines.append(
                    f"- {source_value} -> {canonical_value}"
                )

        if mapping_lines:
            mappings_text = (
                "\n\nUse these approved enterprise-standard "
                "terms where appropriate:\n"
                + "\n".join(mapping_lines)
            )

    instructions = f"""
You are a professional enterprise document translator.

Translate the supplied document into {target_language}.

Requirements:

1. Preserve the meaning of the original document.
2. Preserve paragraph structure as much as possible.
3. Preserve numbers, dates, percentages and monetary amounts accurately.
4. Do not add information that is not present.
5. Keep company names and proper nouns unchanged unless translation
   is clearly appropriate.
6. Use clear professional business language.
7. Return only the translated document text.
8. Do not include commentary before or after the translation.
{mappings_text}
"""

    response = client.responses.create(
        model="gpt-5",
        instructions=instructions,
        input=document_text,
    )

    return response.output_text.strip()


def create_docx_bytes(
    translated_text,
    title="Translated Document",
):
    """
    Create a Word document in memory and return its bytes.
    """

    document = Document()

    document.add_heading(
        title,
        level=1,
    )

    paragraphs = translated_text.split("\n")

    for paragraph in paragraphs:
        paragraph = paragraph.strip()

        if paragraph:
            document.add_paragraph(
                paragraph
            )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output.getvalue()